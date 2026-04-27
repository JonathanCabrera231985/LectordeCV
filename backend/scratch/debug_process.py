import sys
import os
import json
import traceback
import google.generativeai as genai
import pypdf
import docx
from pydantic import BaseModel, Field
from typing import List, Optional

# --- Models from main.py ---
class Universidad(BaseModel):
    institucion: str = "No especificado"
    titulo: str = "No especificado"
    fechas: str = "No especificado"

class Postgrado(BaseModel):
    institucion: str = "No especificado"
    titulo: str = "No especificado"
    fechas: str = "No especificado"

class Educacion(BaseModel):
    universidad: Universidad = Field(default_factory=Universidad)
    postgrado: Postgrado = Field(default_factory=Postgrado)

class Certificacion(BaseModel):
    nombre: str = "No especificado"
    institucion: str = "No especificado"
    horas: str = "No especificado"
    fecha: str = "No especificado"

class Logro(BaseModel):
    nombre: str = "No especificado"
    descripcion: str = "No especificado"
    fecha: str = "No especificado"
    herramientas: str = "No especificado"

class Experiencia(BaseModel):
    empresa: str = "No especificado"
    fecha_ingreso: str = "No especificado"
    fecha_salida: str = "No especificado"
    cargo: str = "No especificado"
    funciones: str = "No especificado"

class Autoevaluacion(BaseModel):
    gestion_proyectos: str = "N/E"
    mitigacion_riesgos: str = "N/E"
    agilismo: str = "N/E"
    cloud: str = "N/E"
    ingenieria_procesos: str = "N/E"
    ms_project: str = "N/E"
    jira: str = "N/E"
    planner: str = "N/E"
    scrum: str = "N/E"
    bpm: str = "N/E"
    power_bi: str = "N/E"
    crm: str = "N/E"
    salesforce: str = "N/E"
    workflow: str = "N/E"
    automatizacion: str = "N/E"
    ibm_filenet: str = "N/E"
    stakeholders: str = "N/E"

class CVData(BaseModel):
    nombre: str = "No especificado"
    pais: str = "No especificado"
    anios_experiencia: str = "No especificado"
    educacion: Educacion = Field(default_factory=Educacion)
    certificaciones: List[Certificacion] = Field(default_factory=list)
    logros: List[Logro] = Field(default_factory=list)
    experiencia: List[Experiencia] = Field(default_factory=list)
    autoevaluacion: Autoevaluacion = Field(default_factory=Autoevaluacion)

# --- Functions ---

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    try:
        reader = pypdf.PdfReader(pdf_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print("Error extracting PDF:", e)
    return text

def safe_set_cell(table, row_idx, col_idx, text):
    try:
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            if col_idx < len(row.cells):
                row.cells[col_idx].text = str(text) if text is not None else "No especificado"
    except Exception as e:
        print(f"Error safe_set_cell ({row_idx}, {col_idx}): {e}")

def fill_docx_template(template_path: str, data: CVData, output_path: str):
    doc = docx.Document(template_path)
    is_new_template = len(doc.tables) >= 13
    print(f"Template has {len(doc.tables)} tables. Is new template: {is_new_template}")

    # 1. Personal
    if len(doc.tables) > 0:
        t = doc.tables[0]
        safe_set_cell(t, 0, 1, data.nombre)
        safe_set_cell(t, 1, 1, data.pais)
    
    # 2. Experience years
    if len(doc.tables) > 1:
        safe_set_cell(doc.tables[1], 0, 1, data.anios_experiencia)

    # 3. Education
    if len(doc.tables) > 2:
        edu = data.educacion
        univ = edu.universidad
        t = doc.tables[2]
        safe_set_cell(t, 1, 1, univ.institucion)
        safe_set_cell(t, 1, 2, univ.titulo)
        safe_set_cell(t, 1, 3, univ.fechas)
        post = edu.postgrado
        safe_set_cell(t, 2, 1, post.institucion)
        safe_set_cell(t, 2, 2, post.titulo)
        safe_set_cell(t, 2, 3, post.fechas)

    # 4. Certs
    if len(doc.tables) > 3:
        certs = data.certificaciones
        table = doc.tables[3]
        for i, cert in enumerate(certs):
            row_idx = i + 1
            if row_idx >= len(table.rows):
                table.add_row()
            safe_set_cell(table, row_idx, 0, cert.nombre)
            safe_set_cell(table, row_idx, 1, cert.institucion)
            date_col = 3 if len(table.rows[0].cells) >= 4 else 2
            safe_set_cell(table, row_idx, date_col, cert.fecha)

    # 5. Logros
    if len(doc.tables) > 4:
        logros = data.logros
        table = doc.tables[4]
        for i, ach in enumerate(logros):
            row_idx = i + 1
            if row_idx >= len(table.rows):
                table.add_row()
            safe_set_cell(table, row_idx, 0, ach.nombre)
            safe_set_cell(table, row_idx, 1, ach.descripcion)
            safe_set_cell(table, row_idx, 2, ach.fecha)
            safe_set_cell(table, row_idx, 3, ach.herramientas)

    # 6. Experience
    def fill_exp(table, emp_data):
        safe_set_cell(table, 0, 0, f"Nombre de la Institución: {emp_data.empresa}")
        safe_set_cell(table, 0, 1, f"Fecha de ingreso: {emp_data.fecha_ingreso}")
        safe_set_cell(table, 0, 2, f"Fecha de finalización del contrato: {emp_data.fecha_salida}")
        safe_set_cell(table, 1, 0, f"Posición: {emp_data.cargo}")
        safe_set_cell(table, 2, 0, f"Principales funciones: {emp_data.funciones}")

    exps = data.experiencia
    start_exp_idx = 5
    for i, exp in enumerate(exps):
        table_idx = start_exp_idx + i
        if table_idx < len(doc.tables):
            if table_idx < (10 if is_new_template else 7):
                fill_exp(doc.tables[table_idx], exp)

    # 7. Autoevaluacion
    auto_idx_start = 10 if is_new_template else 7
    auto = data.autoevaluacion
    if len(doc.tables) > auto_idx_start:
        t = doc.tables[auto_idx_start]
        safe_set_cell(t, 1, 0, auto.gestion_proyectos)
        safe_set_cell(t, 1, 1, auto.mitigacion_riesgos)
        safe_set_cell(t, 1, 2, auto.agilismo)
        safe_set_cell(t, 1, 3, auto.cloud)
        safe_set_cell(t, 1, 4, auto.ingenieria_procesos)

    if len(doc.tables) > auto_idx_start + 1:
        t = doc.tables[auto_idx_start + 1]
        safe_set_cell(t, 1, 0, auto.ms_project)
        safe_set_cell(t, 1, 1, auto.jira)
        safe_set_cell(t, 1, 2, auto.planner)
        safe_set_cell(t, 1, 3, auto.scrum)
        safe_set_cell(t, 1, 4, auto.bpm)
        safe_set_cell(t, 1, 5, auto.power_bi)
        safe_set_cell(t, 1, 6, auto.crm)

    if len(doc.tables) > auto_idx_start + 2:
        t = doc.tables[auto_idx_start + 2]
        safe_set_cell(t, 1, 0, auto.salesforce)
        safe_set_cell(t, 1, 1, f"{auto.workflow} / {auto.automatizacion}")
        safe_set_cell(t, 1, 2, auto.ibm_filenet)
        safe_set_cell(t, 1, 3, auto.stakeholders)

    doc.save(output_path)
    print(f"Saved to {output_path}")

PROMPT_JSON = """
Actúa como un asistente técnico de reclutamiento especializado en la automatización de documentos para SIPECOM S.A.
Extrae la información del siguiente CV y devuelve EXCLUSIVAMENTE un JSON válido sin markdown (`{}`).
Si un dato no existe, coloca "No especificado".
Extructura JSON esperada:
{
  "nombre": "Nombre completo",
  "pais": "Pais de residencia",
  "anios_experiencia": "X años",
  "educacion": {
    "universidad": {"institucion": "...", "titulo": "...", "fechas": "..."},
    "postgrado": {"institucion": "...", "titulo": "...", "fechas": "..."}
  },
  "certificaciones": [
    {"nombre": "...", "institucion": "...", "horas": "...", "fecha": "..."}
  ],
  "logros": [
    {"nombre": "...", "descripcion": "...", "fecha": "...", "herramientas": "..."}
  ],
  "experiencia": [
    {"empresa": "...", "fecha_ingreso": "...", "fecha_salida": "...", "cargo": "...", "funciones": "..."}
  ],
  "autoevaluacion": {
    "gestion_proyectos": "8",
    "mitigacion_riesgos": "7",
    "agilismo": "N/E",
    "cloud": "N/E",
    "ingenieria_procesos": "8",
    "ms_project": "N/E",
    "jira": "N/E",
    "planner": "N/E",
    "scrum": "N/E",
    "bpm": "N/E",
    "power_bi": "8",
    "crm": "N/E",
    "salesforce": "N/E",
    "workflow": "N/E",
    "automatizacion": "7",
    "ibm_filenet": "N/E",
    "stakeholders": "8"
  }
}

CV A ANALIZAR:
"""

# --- Execution ---

API_KEY = "AIzaSyDrY5qC3049u4n2hgC9U4i2KyyN1UjiXTk" # Extracted from .env
CV_FILE = r"c:\Users\jcabrera\TalentoWEB\Antonio Fernandez cv.pdf"
TEMPLATE_FILE = r"c:\Users\jcabrera\TalentoWEB\SDE_Project Manager Plantilla.docx"
OUTPUT_FILE = r"c:\Users\jcabrera\TalentoWEB\backend\scratch\debug_output.docx"

try:
    print("Step 1: Extracting text...")
    text = extract_text_from_pdf(CV_FILE)
    print(f"Extracted {len(text)} characters.")

    print("Step 2: Calling Gemini...")
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel('gemini-flash-latest', generation_config={"response_mime_type": "application/json"})
    response = model.generate_content(PROMPT_JSON + text)
    raw_json = response.text
    print("Gemini response received.")

    print("Step 3: Validating data...")
    parsed_data = CVData.model_validate_json(raw_json)
    print(f"Data validated for: {parsed_data.nombre}")

    print("Step 4: Filling template...")
    fill_docx_template(TEMPLATE_FILE, parsed_data, OUTPUT_FILE)
    print("Process completed successfully.")

except Exception as e:
    print("\n--- ERROR DURING PROCESS ---")
    print(traceback.format_exc())
