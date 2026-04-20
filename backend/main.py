import sys
import os

# Fix para PyInstaller --noconsole (donde stdout/stderr son None)
if sys.stdout is None:
    sys.stdout = open(os.devnull, "w")
if sys.stderr is None:
    sys.stderr = open(os.devnull, "w")

import json
import traceback
import webbrowser
import threading
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
import google.generativeai as genai
import pypdf
import docx
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import List, Optional
from google.api_core import exceptions
import asyncio

# Configuración de logs para evitar errores en modo sin consola
LOG_CONFIG = {
    "version": 1,
    "disable_existing_loggers": False,
    "formatters": {
        "default": {
            "()": "uvicorn.logging.DefaultFormatter",
            "fmt": "%(levelprefix)s %(message)s",
            "use_colors": False,
        },
        "access": {
            "()": "uvicorn.logging.AccessFormatter",
            "fmt": '%(levelprefix)s %(client_addr)s - "%(request_line)s" %(status_code)s',
            "use_colors": False,
        },
    },
    "handlers": {
        "default": {
            "formatter": "default",
            "class": "logging.StreamHandler",
            "stream": "ext://sys.stderr",
        },
        "access": {
            "formatter": "access",
            "class": "logging.StreamHandler",
            "stream": "ext://sys.stdout",
        },
    },
    "loggers": {
        "uvicorn": {"handlers": ["default"], "level": "INFO"},
        "uvicorn.error": {"level": "INFO"},
        "uvicorn.access": {"handlers": ["access"], "level": "INFO", "propagate": False},
    },
}

# Configuración de rutas para PyInstaller
def get_resource_path(relative_path):
    """ Obtiene la ruta absoluta para recursos, compatible con PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def get_exe_dir():
    """ Obtiene el directorio donde reside el ejecutable o el script """
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def load_config():
    """ Carga la API Key desde un archivo externo config.txt con soporte para diferentes codificaciones """
    config_path = os.path.join(get_exe_dir(), "config.txt")
    if os.path.exists(config_path):
        try:
            # utf-8-sig maneja archivos con o sin BOM (común en Windows/PowerShell)
            with open(config_path, "r", encoding="utf-8-sig") as f:
                for line in f:
                    if "GEMINI_API_KEY=" in line:
                        return line.split("=")[1].strip()
        except Exception as e:
            print(f"Error leyendo config.txt: {e}")
    return None


class Universidad(BaseModel):
    institucion: str = "No especificado"
    titulo: str = "No especificado"
    fechas: str = "No especificado"

class Postgrado(BaseModel):
    institucion: str = "No especificado"
    titulo: str = "No especificado"
    fechas: str = "No especificado"

class Educacion(BaseModel):
    universidad: Universidad = Field(default_factory=Universidad)
    postgrado: Postgrado = Field(default_factory=Postgrado)

class Certificacion(BaseModel):
    nombre: str = "No especificado"
    institucion: str = "No especificado"
    horas: str = "No especificado"
    fecha: str = "No especificado"

class Logro(BaseModel):
    nombre: str = "No especificado"
    descripcion: str = "No especificado"
    fecha: str = "No especificado"
    herramientas: str = "No especificado"

class Experiencia(BaseModel):
    empresa: str = "No especificado"
    fecha_ingreso: str = "No especificado"
    fecha_salida: str = "No especificado"
    cargo: str = "No especificado"
    funciones: str = "No especificado"

class Autoevaluacion(BaseModel):
    gestion_proyectos: str = "N/E"
    mitigacion_riesgos: str = "N/E"
    agilismo: str = "N/E"
    cloud: str = "N/E"
    ingenieria_procesos: str = "N/E"
    ms_project: str = "N/E"
    jira: str = "N/E"
    planner: str = "N/E"
    scrum: str = "N/E"
    bpm: str = "N/E"
    power_bi: str = "N/E"
    crm: str = "N/E"
    salesforce: str = "N/E"
    workflow: str = "N/E"
    automatizacion: str = "N/E"
    ibm_filenet: str = "N/E"
    stakeholders: str = "N/E"

class CVData(BaseModel):
    nombre: str = "No especificado"
    pais: str = "No especificado"
    anios_experiencia: str = "No especificado"
    educacion: Educacion = Field(default_factory=Educacion)
    certificaciones: List[Certificacion] = Field(default_factory=list)
    logros: List[Logro] = Field(default_factory=list)
    experiencia: List[Experiencia] = Field(default_factory=list)
    autoevaluacion: Autoevaluacion = Field(default_factory=Autoevaluacion)

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    try:
        reader = pypdf.PdfReader(pdf_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print("Error extracting PDF:", e)
    return text

def safe_set_cell(table, row_idx, col_idx, text):
    """ Escribe en una celda de forma segura, verificando límites """
    try:
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            if col_idx < len(row.cells):
                row.cells[col_idx].text = str(text) if text is not None else "No especificado"
    except Exception as e:
        print(f"Error safe_set_cell (table_idx?, {row_idx}, {col_idx}): {e}")

def fill_docx_template(template_path: str, data: CVData, output_path: str):
    doc = docx.Document(template_path)
    try:
        # Detectar si es la plantilla nueva (13 tablas) o la antigua
        is_new_template = len(doc.tables) >= 13

        # 1. Información Personal
        if len(doc.tables) > 0:
            t = doc.tables[0]
            safe_set_cell(t, 0, 1, data.nombre)
            safe_set_cell(t, 1, 1, data.pais)
        
        # 2. Información Complementaria (Años de experiencia)
        if len(doc.tables) > 1:
            safe_set_cell(doc.tables[1], 0, 1, data.anios_experiencia)

        # 3. Educación
        if len(doc.tables) > 2:
            edu = data.educacion
            univ = edu.university if hasattr(edu, "university") else edu.universidad
            t = doc.tables[2]
            safe_set_cell(t, 1, 1, univ.institucion)
            safe_set_cell(t, 1, 2, univ.titulo)
            safe_set_cell(t, 1, 3, univ.fechas)
            
            post = edu.postgrado
            safe_set_cell(t, 2, 1, post.institucion)
            safe_set_cell(t, 2, 2, post.titulo)
            safe_set_cell(t, 2, 3, post.fechas)

        # 4. Certificaciones
        if len(doc.tables) > 3:
            certs = data.certificaciones
            table = doc.tables[3]
            for i, cert in enumerate(certs):
                row_idx = i + 1
                if row_idx >= len(table.rows):
                    table.add_row()
                
                safe_set_cell(table, row_idx, 0, cert.nombre)
                safe_set_cell(table, row_idx, 1, cert.institucion)
                # Si la tabla tiene 4 columnas, ponemos fecha en la última. 
                # Si tiene 3, la ponemos en la última disponible (la 2).
                date_col = 3 if len(table.rows[0].cells) >= 4 else 2
                safe_set_cell(table, row_idx, date_col, cert.fecha)

        # 5. Logros
        if len(doc.tables) > 4:
            logros = data.logros
            table = doc.tables[4]
            for i, ach in enumerate(logros):
                row_idx = i + 1
                if row_idx >= len(table.rows):
                    table.add_row()
                safe_set_cell(table, row_idx, 0, ach.nombre)
                safe_set_cell(table, row_idx, 1, ach.descripcion)
                safe_set_cell(table, row_idx, 2, ach.fecha)
                safe_set_cell(table, row_idx, 3, ach.herramientas)

        # 6. Experiencia Laboral
        def fill_exp(table, emp_data):
            try:
                safe_set_cell(table, 0, 0, f"Nombre de la Institución: {emp_data.empresa}")
                safe_set_cell(table, 0, 1, f"Fecha de ingreso: {emp_data.fecha_ingreso}")
                safe_set_cell(table, 0, 2, f"Fecha de finalización del contrato: {emp_data.fecha_salida}")
                safe_set_cell(table, 1, 0, f"Posición: {emp_data.cargo}")
                safe_set_cell(table, 2, 0, f"Principales funciones: {emp_data.funciones}")
            except Exception as e:
                print("Error filling exp row:", e)
                
        exps = data.experiencia
        # En la nueva plantilla hay tablas para hasta 5 experiencias (5 a 9)
        start_exp_idx = 5
        for i, exp in enumerate(exps):
            table_idx = start_exp_idx + i
            if table_idx < len(doc.tables):
                # Validar que sea una tabla de experiencia (máximo hasta antes de Autoevaluación)
                if table_idx < (10 if is_new_template else 7):
                    fill_exp(doc.tables[table_idx], exp)

        # 7. Autoevaluación
        # En nueva plantilla son tablas 10, 11, 12. En la vieja 7, 8, 9.
        auto_idx_start = 10 if is_new_template else 7
        
        auto = data.autoevaluacion
        if len(doc.tables) > auto_idx_start:
            t = doc.tables[auto_idx_start]
            safe_set_cell(t, 1, 0, auto.gestion_proyectos)
            safe_set_cell(t, 1, 1, auto.mitigacion_riesgos)
            safe_set_cell(t, 1, 2, auto.agilismo)
            safe_set_cell(t, 1, 3, auto.cloud)
            safe_set_cell(t, 1, 4, auto.ingenieria_procesos)

        if len(doc.tables) > auto_idx_start + 1:
            t = doc.tables[auto_idx_start + 1]
            safe_set_cell(t, 1, 0, auto.ms_project)
            safe_set_cell(t, 1, 1, auto.jira)
            safe_set_cell(t, 1, 2, auto.planner)
            safe_set_cell(t, 1, 3, auto.scrum)
            safe_set_cell(t, 1, 4, auto.bpm)
            safe_set_cell(t, 1, 5, auto.power_bi)
            safe_set_cell(t, 1, 6, auto.crm)

        if len(doc.tables) > auto_idx_start + 2:
            t = doc.tables[auto_idx_start + 2]
            safe_set_cell(t, 1, 0, auto.salesforce)
            safe_set_cell(t, 1, 1, f"{auto.workflow} / {auto.automatizacion}")
            safe_set_cell(t, 1, 2, auto.ibm_filenet)
            safe_set_cell(t, 1, 3, auto.stakeholders)

        doc.save(output_path)
    except Exception as e:
        print("Error filling template:", traceback.format_exc())
        raise Exception("Error procesando docx: " + str(e))

PROMPT_JSON = """
Actúa como un asistente técnico de reclutamiento especializado en la automatización de documentos para SIPECOM S.A.
Extrae la información del siguiente CV y devuelve EXCLUSIVAMENTE un JSON válido sin markdown (`{}`).
Si un dato no existe, coloca "No especificado".
Extructura JSON esperada:
{
  "nombre": "Nombre completo",
  "pais": "Pais de residencia",
  "anios_experiencia": "X años",
  "educacion": {
    "universidad": {"institucion": "...", "titulo": "...", "fechas": "..."},
    "postgrado": {"institucion": "...", "titulo": "...", "fechas": "..."}
  },
  "certificaciones": [
    {"nombre": "...", "institucion": "...", "horas": "...", "fecha": "..."}
  ],
  "logros": [
    {"nombre": "...", "descripcion": "...", "fecha": "...", "herramientas": "..."}
  ],
  "experiencia": [
    {"empresa": "...", "fecha_ingreso": "...", "fecha_salida": "...", "cargo": "...", "funciones": "..."}
  ],
  "autoevaluacion": {
    "gestion_proyectos": "8",
    "mitigacion_riesgos": "7",
    "agilismo": "N/E",
    "cloud": "N/E",
    "ingenieria_procesos": "8",
    "ms_project": "N/E",
    "jira": "N/E",
    "planner": "N/E",
    "scrum": "N/E",
    "bpm": "N/E",
    "power_bi": "8",
    "crm": "N/E",
    "salesforce": "N/E",
    "workflow": "N/E",
    "automatizacion": "7",
    "ibm_filenet": "N/E",
    "stakeholders": "8"
  }
}

CV A ANALIZAR:
"""

async def call_gemini_with_retry(model: genai.GenerativeModel, prompt: str, max_retries: int = 3):
    """ Llama a Gemini con reintentos exponenciales para manejar límites de cuota """
    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt)
            # Validamos que la respuesta tenga texto y no sea una lista vacía de candidatos
            if not response.candidates or not response.text:
                raise Exception("Respuesta vacía de Gemini")
            return response.text
        except exceptions.ResourceExhausted as e:
            if attempt == max_retries - 1:
                raise e
            wait_time = (2 ** attempt) * 10 + 5 # 15s, 25s, ...
            print(f"Cuota excedida. Reintentando en {wait_time}s... (Intento {attempt+1}/{max_retries})")
            await asyncio.sleep(wait_time)
        except Exception as e:
            if attempt == max_retries - 1:
                raise e
            print(f"Error inesperado ({str(e)}). Reintentando en 5s...")
            await asyncio.sleep(5)
    return None

@app.post("/api/generate")
async def generate_cv(
    cv: UploadFile = File(...),
    template: UploadFile = File(...),
    api_key: str = Form(None)
):
    try:
        # Prioridad de API KEY: Formulario > config.txt > .env
        final_api_key = api_key if api_key and api_key.strip() else load_config()
        if not final_api_key:
            final_api_key = os.getenv("GEMINI_API_KEY")

        if not final_api_key:
            raise HTTPException(status_code=400, detail="API Key no encontrada. Por favor, configúrala en el archivo config.txt")

        # Carpeta temp local de ejecución
        base_temp = os.path.join(get_exe_dir(), "temp")
        os.makedirs(base_temp, exist_ok=True)
        
        cv_path = os.path.join(base_temp, cv.filename)
        with open(cv_path, "wb") as f:
            f.write(await cv.read())
            
        template_path = os.path.join(base_temp, template.filename)
        with open(template_path, "wb") as f:
            f.write(await template.read())

        # Extract Text
        cv_text = extract_text_from_pdf(cv_path)

        # Call Gemini with Robustness
        genai.configure(api_key=final_api_key)
        model = genai.GenerativeModel('gemini-flash-latest', generation_config={"response_mime_type": "application/json"})
        
        raw_json = await call_gemini_with_retry(model, PROMPT_JSON + cv_text)
        if not raw_json:
            raise HTTPException(status_code=500, detail="No se pudo obtener una respuesta válida de la IA tras varios intentos.")
        
        # Validation with Pydantic
        try:
            parsed_data = CVData.model_validate_json(raw_json)
        except Exception as ve:
            print(f"Validation Error: {ve}")
            # Fallback a parseo simple si falla la validación estricta
            parsed_data = CVData(**json.loads(raw_json))

        output_filename = f"SOLICITUD_{parsed_data.nombre.replace(' ', '_')}.docx"
        output_path = os.path.join(base_temp, output_filename)

        # Fill DOCX
        fill_docx_template(template_path, parsed_data, output_path)

        return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=output_filename)

    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

# Servir Frontend
frontend_path = get_resource_path("frontend/dist")
if os.path.exists(frontend_path):
    app.mount("/", StaticFiles(directory=frontend_path, html=True), name="frontend")

def open_browser():
    webbrowser.open("http://localhost:8000")

if __name__ == "__main__":
    import uvicorn
    # Iniciar navegador en un hilo separado
    threading.Timer(1.5, open_browser).start()
    # Usamos LOG_CONFIG explícito con use_colors=False para evitar el error isatty
    uvicorn.run(app, host="0.0.0.0", port=8000, log_config=LOG_CONFIG)
